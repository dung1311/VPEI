from docxtpl import DocxTemplate
from docx import Document
from docx.shared import RGBColor, Pt
from openpyxl import load_workbook
import io
import os
import math
from copy import copy
from calendar import monthrange
from datetime import date, datetime, time
from typing import Any, Optional, Set, Tuple, List
from sqlalchemy.orm import Session
from sqlalchemy import extract, or_, and_

from core.vn_format import journey_type_label_vn

# [UPDATED] Import Device thay cho DeviceCategory
from models.device import ActivityData, Device 
from models.electrical_item import ElectricalItem
from models.ship import Ship
from models.container import Container
from models.harbor_craft import HarborCraft
from services.ship_service import calculate_ship_co2
from services.scope3_period_service import compute_scope3_period

def _report_months(month: Optional[int], quarter: Optional[int]) -> Optional[Set[int]]:
    if month is not None:
        return {month}
    if quarter is not None:
        q = int(quarter)
        return set(range((q - 1) * 3 + 1, q * 3 + 1))
    return None

def _parse_period_value_date(raw: Optional[str]) -> Optional[date]:
    s = (raw or "").strip()
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _scope2_period_display(raw: Optional[str]) -> str:
    """period_value trong DB thường là yyyy-mm-dd — báo cáo chỉ hiển thị ngày dd/mm/yyyy."""
    d = _parse_period_value_date(raw)
    if d is None:
        t = (raw or "").strip()
        return t if t else "N/A"
    return d.strftime("%d/%m/%Y")


def _dt_sort_key(dt: Optional[datetime]) -> datetime:
    """Dùng khi sắp xếp giảm dần: bản ghi không có ngày xếp cuối."""
    return dt if dt is not None else datetime.min


def _sort_rows_by_dt(
    parts: List[Tuple[datetime, List[Any]]],
    reverse: bool = True,
) -> List[List[Any]]:
    parts_sorted = sorted(parts, key=lambda x: x[0], reverse=reverse)
    return [[i + 1, *row] for i, (_, row) in enumerate(parts_sorted)]


def _year_built_sort_key(yb: Any) -> int:
    """Năm đóng tàu để sort giảm dần (mới nhất trên); thiếu/không hợp lệ → 0 (xếp cuối)."""
    try:
        y = int(yb)
        return y if y > 0 else 0
    except (TypeError, ValueError):
        return 0


def _sort_rows_by_year_built(
    parts: List[Tuple[int, List[Any]]],
    reverse: bool = True,
) -> List[List[Any]]:
    parts_sorted = sorted(parts, key=lambda x: x[0], reverse=reverse)
    return [[i + 1, *row] for i, (_, row) in enumerate(parts_sorted)]


def _year_month_pairs_in_range(d_start: date, d_end: date) -> List[Tuple[int, int]]:
    out: List[Tuple[int, int]] = []
    cur = date(d_start.year, d_start.month, 1)
    while cur <= d_end:
        last = date(cur.year, cur.month, monthrange(cur.year, cur.month)[1])
        if last >= d_start and cur <= d_end:
            out.append((cur.year, cur.month))
        if cur.month == 12:
            cur = date(cur.year + 1, 1, 1)
        else:
            cur = date(cur.year, cur.month + 1, 1)
    return out


class ReportService:
    @staticmethod
    def set_cell(cell, text):
        cell.text = "" 
        run = cell.paragraphs[0].add_run(str(text))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        run.font.highlight_color = None

    @staticmethod
    def generate_vpei_final_report(
        db: Session,
        year: int,
        month: Optional[int] = None,
        quarter: Optional[int] = None,
        date_start: Optional[date] = None,
        date_end: Optional[date] = None,
        include_appendix: bool = True,
    ) -> bytes:
        use_range = date_start is not None and date_end is not None
        if use_range:
            if date_start > date_end:
                raise ValueError("date_start phải trước hoặc bằng date_end")
            start_time_s = date_start.strftime("%d/%m/%Y")
            end_time_s = date_end.strftime("%d/%m/%Y")
            mf = None
        else:
            mf = _report_months(month, quarter)
            if month is not None:
                d1 = monthrange(year, month)[1]
                start_time_s = f"01/{month:02d}/{year}"
                end_time_s = f"{d1:02d}/{month:02d}/{year}"
            elif quarter is not None:
                qn = int(quarter)
                sm, em = (qn - 1) * 3 + 1, qn * 3
                d1 = monthrange(year, em)[1]
                start_time_s = f"01/{sm:02d}/{year}"
                end_time_s = f"{d1:02d}/{em:02d}/{year}"
            else:
                start_time_s = f"01/01/{year}"
                end_time_s = f"31/12/{year}"

        # ==========================================
        # 1. Fetch Real Data from Database
        # ==========================================

        # --- Scope 1: ActivityData ---
        # [UPDATED] JOIN với Device thay vì DeviceCategory
        if use_range:
            pairs = _year_month_pairs_in_range(date_start, date_end)
            if pairs:
                q_s1 = db.query(ActivityData).join(Device).filter(
                    or_(*[and_(ActivityData.period_year == py, ActivityData.period_month == pm) for py, pm in pairs])
                )
            else:
                q_s1 = db.query(ActivityData).join(Device).filter(ActivityData.id == -1)
            scope1_activities = q_s1.order_by(ActivityData.record_time.desc()).all()
        else:
            q_s1 = db.query(ActivityData).join(Device).filter(ActivityData.period_year == year)
            if mf is not None:
                q_s1 = q_s1.filter(ActivityData.period_month.in_(mf))
            scope1_activities = q_s1.order_by(ActivityData.record_time.desc()).all()

        s1_co2e = sum(act.total_co2e for act in scope1_activities) if scope1_activities else 0.0
        s1_co2 = s1_co2e
        s1_ch4 = 0.0
        s1_n2o = 0.0

        data_scope1 = []
        for idx, act in enumerate(scope1_activities, 1):
            # [UPDATED] Lấy tên từ bảng Device
            device_name = act.device.name if act.device else "Unknown Device"
            operating_hours = act.operating_hours
            power = act.recorded_power
            co2 = round(act.total_co2e, 2)
            
            # Thay vì "N/A", hiển thị Thời gian ghi nhận để báo cáo chi tiết hơn
            record_time_str = act.record_time.strftime("%d/%m/%Y") if getattr(act, 'record_time', None) else "N/A"
            
            data_scope1.append([idx, device_name, record_time_str, power, operating_hours, co2])

        # --- Scope 2: Electricity ---
        all_elec = db.query(ElectricalItem).all()
        if use_range:
            scope2_items = []
            for e in all_elec:
                dv = _parse_period_value_date(e.period_value)
                if dv is not None and date_start <= dv <= date_end:
                    scope2_items.append(e)
        else:
            scope2_items = [e for e in all_elec if e.period_value and str(year) in e.period_value]

        scope2_items.sort(
            key=lambda e: _parse_period_value_date(e.period_value) or date.min,
            reverse=True,
        )

        s2_co2e = 0.0
        data_scope2 = []
        for idx, item in enumerate(scope2_items, 1):
            e_total = (item.power * 720 * 0.8 * 0.6235 / 1000) if item.power else 0.0
            if item.period_type == 'year':
                e_total *= 12

            s2_co2e += e_total
            data_scope2.append([idx, item.name, _scope2_period_display(item.period_value), f"{item.power} kW", round(e_total, 2)])

        s2_co2 = s2_co2e
        s2_ch4 = 0.0
        s2_n2o = 0.0

        # --- Scope 3: Ships, Containers ---
        if use_range:
            t0 = datetime.combine(date_start, time.min)
            t1 = datetime.combine(date_end, time.max)
            q_sh = db.query(Ship).filter(Ship.start_time.isnot(None), Ship.start_time >= t0, Ship.start_time <= t1)
            q_ct = db.query(Container).filter(Container.start_time.isnot(None), Container.start_time >= t0, Container.start_time <= t1)
            ships = q_sh.order_by(Ship.start_time.desc()).all()
            containers = q_ct.order_by(Container.start_time.desc()).all()
        else:
            q_sh = db.query(Ship).filter(extract("year", Ship.start_time) == year)
            q_ct = db.query(Container).filter(extract("year", Container.start_time) == year)
            if mf is not None:
                mlist = list(mf)
                q_sh = q_sh.filter(extract("month", Ship.start_time).in_(mlist))
                q_ct = q_ct.filter(extract("month", Container.start_time).in_(mlist))
            ships = q_sh.order_by(Ship.start_time.desc()).all()
            containers = q_ct.order_by(Container.start_time.desc()).all()

        if use_range:
            harbor_crafts = (
                db.query(HarborCraft)
                .filter(
                    HarborCraft.record_time.isnot(None),
                    HarborCraft.record_time >= t0,
                    HarborCraft.record_time <= t1,
                )
                .order_by(HarborCraft.record_time.desc())
                .all()
            )
        else:
            q_hc = db.query(HarborCraft).filter(extract("year", HarborCraft.record_time) == year)
            if mf is not None:
                q_hc = q_hc.filter(extract("month", HarborCraft.record_time).in_(list(mf)))
            harbor_crafts = q_hc.order_by(HarborCraft.record_time.desc()).all()

        s3_co2e_ships = 0.0
        parts_tau_bien: List[Tuple[datetime, List[Any]]] = []
        parts_tau_cang: List[Tuple[int, List[Any]]] = []
        for ship in ships:
            co2 = calculate_ship_co2(ship)
            s3_co2e_ships += co2

            start_str = ship.start_time.strftime("%d/%m/%Y %H:%M") if ship.start_time else ""
            end_str = ship.end_time.strftime("%d/%m/%Y %H:%M") if ship.end_time else ""
            st = _dt_sort_key(ship.start_time)

            if "lai dắt" in ship.name.lower() or "hoa tiêu" in ship.name.lower() or ship.ship_type == "Tàu cảng":
                yk = _year_built_sort_key(ship.year_built)
                parts_tau_cang.append(
                    (yk, [ship.name, ship.year_built, ship.P_main, ship.P_main, round(co2, 2)])
                )
            else:
                parts_tau_bien.append(
                    (
                        st,
                        [
                            ship.name,
                            start_str,
                            end_str,
                            ship.deadweight_tonnage,
                            ship.P_main,
                            round(co2, 2),
                        ],
                    )
                )

        s3_co2e_containers = 0.0
        data_xe_cont = []
        for idx, con in enumerate(containers, 1):
            co2 = con.e_total if con.e_total else 0.0
            s3_co2e_containers += co2

            start_str = con.start_time.strftime("%d/%m/%Y %H:%M") if con.start_time else ""
            end_str = con.end_time.strftime("%d/%m/%Y %H:%M") if con.end_time else ""
            data_xe_cont.append(
                [idx, con.license_plate, start_str, end_str, journey_type_label_vn(con.journey_type), round(co2, 2)]
            )

        for h in harbor_crafts:
            co2_h = float(h.e_total or 0.0)
            eng = h.engine_type.value if hasattr(h.engine_type, "value") else str(h.engine_type)
            eng_l = str(eng).lower()
            if eng_l == "aux":
                p_main, p_aux = "-", round(h.power, 2)
            else:
                p_main, p_aux = round(h.power, 2), "-"
            yk = _year_built_sort_key(h.year_built)
            parts_tau_cang.append((yk, [h.device_name, h.year_built, p_main, p_aux, round(co2_h, 2)]))

        data_tau_bien = _sort_rows_by_dt(parts_tau_bien)
        data_tau_cang = _sort_rows_by_year_built(parts_tau_cang)

        s3_co2e_other = 0.0
        xe_may_count = 0
        xe_may_co2 = 0.0
        oto_count = 0
        oto_co2 = 0.0

        # Tổng Scope 3 khớp dashboard / trang Scope 3: gồm tàu biển + xe container + tàu cảng (harbor craft).
        s3_co2e_harbor = sum(float(h.e_total or 0.0) for h in harbor_crafts)
        if use_range:
            s3_co2e = s3_co2e_ships + s3_co2e_containers + s3_co2e_harbor + s3_co2e_other
        else:
            s3p = compute_scope3_period(db, year, month, quarter)
            s3_co2e = float(s3p.get("total_co2e") or 0.0)
        s3_co2 = s3_co2e
        s3_ch4 = 0.0
        s3_n2o = 0.0

        sum_co2e = s1_co2e + s2_co2e + s3_co2e
        
        if sum_co2e == 0:
            sum_co2e = 1.0

        s1_w = f"{round((s1_co2e / sum_co2e) * 100, 2)}%"
        s2_w = f"{round((s2_co2e / sum_co2e) * 100, 2)}%"
        s3_w = f"{round((s3_co2e / sum_co2e) * 100, 2)}%"

        context = {
            'company_name': 'Công ty TNHH Cảng Nam Đình Vũ',
            'start_time': start_time_s,
            'end_time': end_time_s,
            'seaport_name': 'Nam Đình Vũ - Hải Phòng',
            's1_co2': round(s1_co2, 2), 's1_ch4': round(s1_ch4, 2), 's1_n2o': round(s1_n2o, 2), 's1_co2e': round(s1_co2e, 2), 's1_w': s1_w,
            's2_co2': round(s2_co2, 2), 's2_ch4': round(s2_ch4, 2), 's2_n2o': round(s2_n2o, 2), 's2_co2e': round(s2_co2e, 2), 's2_w': s2_w,
            's3_co2': round(s3_co2, 2), 's3_ch4': round(s3_ch4, 2), 's3_n2o': round(s3_n2o, 2), 's3_co2e': round(s3_co2e, 2), 's3_w': s3_w,
            'sum_co2e': round(sum_co2e, 2) if sum_co2e != 1.0 else 0.0
        }

        template_path = "./static/VPEI_report_apd.docx" if include_appendix else "./static/VPEI_report.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template {template_path} not found")

        doc = DocxTemplate(template_path)
        doc.render(context)
        
        temp_io = io.BytesIO()
        doc.save(temp_io)
        temp_io.seek(0)

        # ==========================================
        # 2. FILL TABLES USING PYTHON-DOCX
        # ==========================================
        final_doc = Document(temp_io)

        try:
            for item in data_scope1:
                cells = final_doc.tables[2].add_row().cells
                for idx, val in enumerate(item): ReportService.set_cell(cells[idx], val)

            for item in data_scope2:
                cells = final_doc.tables[3].add_row().cells
                for idx, val in enumerate(item): ReportService.set_cell(cells[idx], val)

            for item in data_tau_bien:
                cells = final_doc.tables[4].add_row().cells
                for idx, val in enumerate(item): ReportService.set_cell(cells[idx], val)

            for item in data_xe_cont:
                cells = final_doc.tables[5].add_row().cells
                for idx, val in enumerate(item): ReportService.set_cell(cells[idx], val)

            for item in data_tau_cang:
                cells = final_doc.tables[6].add_row().cells
                for idx, val in enumerate(item): ReportService.set_cell(cells[idx], val)

            if len(final_doc.tables) > 7:
                table_phuong_tien = final_doc.tables[7]
                ReportService.set_cell(table_phuong_tien.rows[1].cells[1], xe_may_count)
                ReportService.set_cell(table_phuong_tien.rows[1].cells[3], round(xe_may_co2, 2))
                ReportService.set_cell(table_phuong_tien.rows[2].cells[1], oto_count)
                ReportService.set_cell(table_phuong_tien.rows[2].cells[3], round(oto_co2, 2))

        except IndexError as e:
            print(f"Lỗi cấu trúc bảng: {e}")

        final_output_io = io.BytesIO()
        final_doc.save(final_output_io)
        final_output_io.seek(0)
        return final_output_io.read()

    @staticmethod
    def _enum_value(v: Any) -> Any:
        return v.value if hasattr(v, "value") else v

    @staticmethod
    def _copy_row_style(ws, src_row: int, dst_row: int, max_cols: int) -> None:
        if src_row < 1 or src_row > ws.max_row:
            return
        src_dim = ws.row_dimensions.get(src_row)
        if src_dim and src_dim.height is not None:
            ws.row_dimensions[dst_row].height = src_dim.height

        for c in range(1, max_cols + 1):
            src = ws.cell(src_row, c)
            dst = ws.cell(dst_row, c)
            dst._style = copy(src._style)
            dst.number_format = src.number_format
            dst.font = copy(src.font)
            dst.fill = copy(src.fill)
            dst.border = copy(src.border)
            dst.alignment = copy(src.alignment)
            dst.protection = copy(src.protection)

    @staticmethod
    def _write_rows_to_sheet(ws, start_row: int, rows: List[List[Any]], max_cols: int, style_row: Optional[int] = None) -> None:
        style_row = style_row or start_row
        old_max_row = ws.max_row
        existing_slots = max(old_max_row - start_row + 1, 0)
        extra_rows = len(rows) - existing_slots

        if extra_rows > 0:
            insert_at = old_max_row + 1 if old_max_row >= start_row else start_row
            ws.insert_rows(insert_at, extra_rows)
            for r in range(insert_at, insert_at + extra_rows):
                ReportService._copy_row_style(ws, style_row, r, max_cols)

        for i, row_vals in enumerate(rows):
            r = start_row + i
            for c in range(1, max_cols + 1):
                ws.cell(r, c).value = row_vals[c - 1] if c - 1 < len(row_vals) else None

        clear_from = start_row + len(rows)
        clear_to = old_max_row
        if clear_from <= clear_to:
            for r in range(clear_from, clear_to + 1):
                for c in range(1, max_cols + 1):
                    ws.cell(r, c).value = None

    @staticmethod
    def generate_vpei_excel_report(
        db: Session,
        year: int,
        month: Optional[int] = None,
        quarter: Optional[int] = None,
        date_start: Optional[date] = None,
        date_end: Optional[date] = None,
    ) -> bytes:
        use_range = date_start is not None and date_end is not None
        if use_range:
            if date_start > date_end:
                raise ValueError("date_start phải trước hoặc bằng date_end")
            start_time_s = date_start.strftime("%d/%m/%Y")
            end_time_s = date_end.strftime("%d/%m/%Y")
            mf = None
        else:
            mf = _report_months(month, quarter)
            if month is not None:
                d1 = monthrange(year, month)[1]
                start_time_s = f"01/{month:02d}/{year}"
                end_time_s = f"{d1:02d}/{month:02d}/{year}"
            elif quarter is not None:
                qn = int(quarter)
                sm, em = (qn - 1) * 3 + 1, qn * 3
                d1 = monthrange(year, em)[1]
                start_time_s = f"01/{sm:02d}/{year}"
                end_time_s = f"{d1:02d}/{em:02d}/{year}"
            else:
                start_time_s = f"01/01/{year}"
                end_time_s = f"31/12/{year}"

        # Scope 1
        if use_range:
            pairs = _year_month_pairs_in_range(date_start, date_end)
            if pairs:
                q_s1 = db.query(ActivityData).join(Device).filter(
                    or_(*[and_(ActivityData.period_year == py, ActivityData.period_month == pm) for py, pm in pairs])
                )
            else:
                q_s1 = db.query(ActivityData).join(Device).filter(ActivityData.id == -1)
            scope1_activities = q_s1.order_by(ActivityData.record_time.desc()).all()
        else:
            q_s1 = db.query(ActivityData).join(Device).filter(ActivityData.period_year == year)
            if mf is not None:
                q_s1 = q_s1.filter(ActivityData.period_month.in_(mf))
            scope1_activities = q_s1.order_by(ActivityData.record_time.desc()).all()

        scope1_rows: List[List[Any]] = []
        for idx, act in enumerate(scope1_activities, 1):
            device_name = act.device.name if act.device else "Unknown Device"
            scope1_rows.append(
                [
                    idx,
                    device_name,
                    "-",
                    round(float(act.recorded_power or 0.0), 2),
                    act.period_month,
                    act.period_year,
                    round(float(act.operating_hours or 0.0), 2),
                    round(float(act.total_co2e or 0.0), 2),
                ]
            )

        # Scope 2
        all_elec = db.query(ElectricalItem).all()
        if use_range:
            scope2_items = []
            for e in all_elec:
                dv = _parse_period_value_date(e.period_value)
                if dv is not None and date_start <= dv <= date_end:
                    scope2_items.append(e)
        else:
            scope2_items = [e for e in all_elec if e.period_value and str(year) in e.period_value]

        scope2_items.sort(
            key=lambda e: _parse_period_value_date(e.period_value) or date.min,
            reverse=True,
        )

        scope2_rows: List[List[Any]] = []
        for item in scope2_items:
            e_total = (item.power * 720 * 0.8 * 0.6235 / 1000) if item.power else 0.0
            if item.period_type == "year":
                e_total *= 12

            dv = _parse_period_value_date(item.period_value)
            month_val: Any = dv.month if dv else ""
            year_val: Any = dv.year if dv else year
            if item.period_type == "year":
                month_val = "Cả năm"

            scope2_rows.append(
                [
                    item.name,
                    month_val,
                    year_val,
                    round(float(item.power or 0.0), 2),
                    0.6235,
                    round(float(e_total), 2),
                ]
            )

        # Scope 3
        if use_range:
            t0 = datetime.combine(date_start, time.min)
            t1 = datetime.combine(date_end, time.max)
            q_sh = db.query(Ship).filter(Ship.start_time.isnot(None), Ship.start_time >= t0, Ship.start_time <= t1)
            q_ct = db.query(Container).filter(Container.start_time.isnot(None), Container.start_time >= t0, Container.start_time <= t1)
            ships = q_sh.order_by(Ship.start_time.desc()).all()
            containers = q_ct.order_by(Container.start_time.desc()).all()
            harbor_crafts = (
                db.query(HarborCraft)
                .filter(
                    HarborCraft.record_time.isnot(None),
                    HarborCraft.record_time >= t0,
                    HarborCraft.record_time <= t1,
                )
                .order_by(HarborCraft.record_time.desc())
                .all()
            )
        else:
            q_sh = db.query(Ship).filter(extract("year", Ship.start_time) == year)
            q_ct = db.query(Container).filter(extract("year", Container.start_time) == year)
            if mf is not None:
                mlist = list(mf)
                q_sh = q_sh.filter(extract("month", Ship.start_time).in_(mlist))
                q_ct = q_ct.filter(extract("month", Container.start_time).in_(mlist))
            ships = q_sh.order_by(Ship.start_time.desc()).all()
            containers = q_ct.order_by(Container.start_time.desc()).all()

            q_hc = db.query(HarborCraft).filter(extract("year", HarborCraft.record_time) == year)
            if mf is not None:
                q_hc = q_hc.filter(extract("month", HarborCraft.record_time).in_(list(mf)))
            harbor_crafts = q_hc.order_by(HarborCraft.record_time.desc()).all()

        ship_rows: List[List[Any]] = []
        for idx, ship in enumerate(ships, 1):
            co2 = round(float(calculate_ship_co2(ship)), 2)
            start_dt = ship.start_time
            end_dt = ship.end_time
            start_str = start_dt.strftime("%d/%m/%Y %H:%M") if start_dt else ""
            end_str = end_dt.strftime("%d/%m/%Y %H:%M") if end_dt else ""
            month_val = start_dt.month if start_dt else ""
            time_in_port = round(float(ship.time_in_port or 0.0), 2)
            ship_type_val = ReportService._enum_value(ship.ship_type) or ""
            valve_type_val = ReportService._enum_value(ship.valve_type) or ""
            man_label = "MAN" if bool(ship.is_man) else "NON-MAN"

            ship_rows.append(
                [
                    idx,                              # A
                    ship.name,                        # B
                    ship_type_val,                    # C
                    round(float(ship.deadweight_tonnage or 0.0), 2),  # D
                    ship.year_built,                  # E
                    month_val,                        # F
                    start_str,                        # G
                    end_str,                          # H
                    time_in_port,                     # I
                    None,                             # J
                    None,                             # K
                    None,                             # L
                    round(float(ship.v_trip or 0.0), 2),      # M
                    round(float(ship.v_maneuver or 0.0), 2),  # N
                    round(float(ship.v_max or 0.0), 2),       # O
                    round(float(ship.P_main or 0.0), 2),      # P
                    round(float(ship.P_aux or 0.0), 2),       # Q
                    valve_type_val,                   # R
                    man_label,                        # S
                    round(float(ship.rpm or 0.0), 2),         # T
                    None,                             # U NOx
                    None,                             # V PM10
                    None,                             # W PM2.5
                    None,                             # X HC
                    None,                             # Y CO
                    co2,                              # Z CO2
                    None,                             # AA SOx
                    None,                             # AB N2O
                    None,                             # AC CH4
                    None,                             # AD E1
                    None,                             # AE E2
                    None,                             # AF E3
                    None,                             # AG E4
                    None,                             # AH E5
                    co2,                              # AI Tổng
                ]
            )

        container_rows: List[List[Any]] = []
        for idx, con in enumerate(containers, 1):
            start_dt = con.start_time
            end_dt = con.end_time
            start_str = start_dt.strftime("%d/%m/%Y %H:%M") if start_dt else ""
            end_str = end_dt.strftime("%d/%m/%Y %H:%M") if end_dt else ""

            duration = con.duration
            if duration is None and start_dt and end_dt:
                duration = (end_dt - start_dt).total_seconds() / 3600.0

            move_time = float(con.time1 or 0.0) + float(con.time2 or 0.0) + float(con.time3 or 0.0)
            total_distance = float(con.distance_1 or 0.0) + float(con.distance_2 or 0.0) + float(con.distance_3 or 0.0)
            velocity_move = con.velocity if con.velocity is not None else (
                (float(con.velocity_1 or 0.0) + float(con.velocity_2 or 0.0) + float(con.velocity_3 or 0.0)) / 3.0
            )

            truck_type = "lạnh" if getattr(con, "is_refrigerated", False) else "thường"

            container_rows.append(
                [
                    idx,
                    con.license_plate,
                    journey_type_label_vn(con.journey_type),
                    "-",
                    truck_type,
                    start_str,
                    end_str,
                    round(float(duration or 0.0), 2),
                    round(move_time, 2),
                    round(float(con.waited_time or 0.0), 2),
                    round(float(velocity_move or 0.0), 2),
                    round(float(con.active_waited_time or 0.0), 2),
                    round(total_distance, 2),
                    round(float(con.e_total or 0.0), 2),
                ]
            )

        harbor_rows: List[List[Any]] = []
        for idx, h in enumerate(harbor_crafts, 1):
            rec = h.record_time
            month_val = rec.month if rec else ""
            year_val = rec.year if rec else ""
            eng = str(ReportService._enum_value(h.engine_type) or "").lower()
            p_main = round(float(h.power or 0.0), 2) if eng != "aux" else "-"

            harbor_rows.append(
                [
                    idx,
                    h.device_name,
                    h.year_built,
                    p_main,
                    month_val,
                    year_val,
                    round(float(h.activity_hours or 0.0), 2),
                    round(float(h.e_total or 0.0), 2),
                ]
            )

        xe_may_count = 0
        xe_may_co2 = 0.0
        oto_count = 0
        oto_co2 = 0.0

        template_path = "./static/Báo-cáo-VPEI-file-excel.xlsx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template {template_path} not found")

        wb = load_workbook(template_path)
        if len(wb.sheetnames) < 7:
            raise ValueError("Template Excel không đúng định dạng 7 sheet")

        ws_general = wb["Báo cáo chung"]
        ws_s1 = wb["Scope 1"]
        ws_s2 = wb["Scope 2"]
        ws_s3_ship = wb["Scope 3 - Tàu biển"]
        ws_s3_container = wb["Scope 3 - Xe contaner"]
        ws_s3_harbor = wb["Scope 3 - Tàu cảng"]
        ws_s3_commute = wb["Scope 3 - Phương tiện đi lại"]

        ws_general["C3"] = "Công ty TNHH Cảng Nam Đình Vũ"
        ws_general["C4"] = f"Tính từ ngày {start_time_s}"
        ws_general["C5"] = f"Tính đến ngày {end_time_s}"
        ws_general["C7"] = "Cảng Nam Đình Vũ"

        ReportService._write_rows_to_sheet(ws_s1, start_row=3, rows=scope1_rows, max_cols=8, style_row=3)
        ReportService._write_rows_to_sheet(ws_s2, start_row=3, rows=scope2_rows, max_cols=6, style_row=3)
        ReportService._write_rows_to_sheet(ws_s3_ship, start_row=6, rows=ship_rows, max_cols=35, style_row=6)
        ReportService._write_rows_to_sheet(ws_s3_container, start_row=7, rows=container_rows, max_cols=14, style_row=7)
        ReportService._write_rows_to_sheet(ws_s3_harbor, start_row=3, rows=harbor_rows, max_cols=8, style_row=3)

        ws_s3_commute["B3"] = xe_may_count
        ws_s3_commute["D3"] = round(xe_may_co2, 2)
        ws_s3_commute["B4"] = oto_count
        ws_s3_commute["D4"] = round(oto_co2, 2)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.read()